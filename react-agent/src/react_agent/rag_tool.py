# RAG (Retrieval-Augmented Generation) 도구
# Chroma DB를 사용한 벡터 검색 및 문서 검색 기능

from __future__ import annotations

import os
import logging
import time
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple, TYPE_CHECKING
from react_agent.cache_manager import get_cache_manager

try:
    from langchain_chroma import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    from rank_bm25 import BM25Okapi
    import numpy as np
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

if TYPE_CHECKING:
    from react_agent.knowledge_graph import Neo4jGraphManager

logger = logging.getLogger(__name__)
perf_logger = logging.getLogger("perf")


class RAGTool:
    """RAG 검색 도구 클래스"""

    def __init__(
        self,
        knowledge_base_path: Optional[str] = None,
        chroma_db_path: Optional[str] = None
    ):
        """
        RAG 도구 초기화

        Args:
            knowledge_base_path: 지식베이스 문서 경로
            chroma_db_path: Chroma DB 저장 경로
        """
        if not RAG_AVAILABLE:
            logger.warning("RAG 라이브러리가 설치되지 않았습니다.")
            self.available = False
            return

        self.available = True
        self._kb_last_modified: Optional[float] = None  # 지식베이스 마지막 수정 시간

        # 경로 설정
        if knowledge_base_path is None:
            knowledge_base_path = os.getenv(
                "KNOWLEDGE_BASE_PATH",
                str(Path(__file__).parent.parent.parent / "knowledge_base")
            )
        if chroma_db_path is None:
            chroma_db_path = os.getenv(
                "CHROMA_DB_PATH",
                str(Path(__file__).parent.parent.parent / "chroma_db")
            )

        self.knowledge_base_path = Path(knowledge_base_path)
        self.chroma_db_path = Path(chroma_db_path)

        # 임베딩 모델 초기화
        try:
            # HF_TOKEN 설정 (있으면 사용, 없으면 무시)
            hf_token = os.environ.get("HF_TOKEN")
            if hf_token:
                os.environ["HUGGINGFACE_HUB_TOKEN"] = hf_token

            # Safetensors 자동 변환 비활성화 (타임아웃 방지)
            os.environ["TRANSFORMERS_OFFLINE"] = "0"  # 온라인 유지
            os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"  # 텔레메트리 비활성화

            self.embeddings = HuggingFaceEmbeddings(
                model_name="dragonkue/BGE-m3-ko",
                model_kwargs={
                    'device': 'cpu',
                    'trust_remote_code': False  # 보안 강화
                },
                encode_kwargs={'normalize_embeddings': True}  # 벡터 정규화 활성화
            )
            logger.info("한국어 임베딩 모델 로드 완료: BGE-m3-ko (1024-dim, 정규화 활성화)")
        except Exception as e:
            logger.warning(f"한국어 임베딩 모델 로드 실패, 기본 모델 사용: {e}")
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                encode_kwargs={'normalize_embeddings': True}  # 벡터 정규화 활성화
            )

        # 텍스트 분할기 (의미적 청킹 전략)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=150,
            length_function=len,
            separators=[
                "\n\n\n", "\n\n", "\n", ". ", "。", "! ", "? ", ".", ", ", "，", " ", ""
            ],
            keep_separator=True
        )

        # 벡터 스토어 (지연 로딩)
        self._vectorstore: Optional[Chroma] = None

        # BM25 인덱스 (지연 로딩)
        self._bm25_index: Optional['BM25Okapi'] = None
        self._bm25_documents: List[Document] = []

        # 그래프 매니저 (지연 로딩)
        self._graph_manager: Optional['Neo4jGraphManager'] = None

        # 지식베이스 변경 감지 초기화
        self._update_kb_modified_time()

        logger.info(f"RAG 도구 초기화 완료: {knowledge_base_path}")

    def _get_kb_modified_time(self) -> Optional[float]:
        """지식베이스 디렉토리의 최신 수정 시간 반환"""
        if not self.knowledge_base_path.exists():
            return None

        try:
            latest_time = 0.0
            for ext in ['.txt', '.md', '.pdf', '.docx']:
                for file_path in self.knowledge_base_path.rglob(f"*{ext}"):
                    mtime = file_path.stat().st_mtime
                    if mtime > latest_time:
                        latest_time = mtime
            return latest_time if latest_time > 0 else None
        except Exception as e:
            logger.error(f"지식베이스 수정 시간 확인 실패: {e}")
            return None

    def _update_kb_modified_time(self):
        """지식베이스 수정 시간 업데이트"""
        self._kb_last_modified = self._get_kb_modified_time()
        if self._kb_last_modified:
            logger.info(f"지식베이스 마지막 수정 시간: {self._kb_last_modified}")

    def _check_kb_changed(self) -> bool:
        """지식베이스가 변경되었는지 확인"""
        current_time = self._get_kb_modified_time()
        if current_time is None or self._kb_last_modified is None:
            return False

        changed = current_time > self._kb_last_modified
        if changed:
            logger.info("지식베이스 변경 감지! 캐시를 클리어합니다.")
            cache_manager = get_cache_manager()
            cache_manager.clear(prefix="rag")
            cache_manager.clear(prefix="llm")
            self._kb_last_modified = current_time

        return changed

    def _extract_keywords_from_text(self, text: str, max_keywords: int = 5) -> List[str]:
        """텍스트에서 핵심 키워드 추출"""
        stopwords = {'은', '는', '이', '가', '을', '를', '의', '에', '로', '와', '과', '도', '만',
                     '하다', '있다', '되다', '않다', '같다', '위해', '대한', '통해', '따라'}

        words = text.split()
        keywords = []

        for word in words:
            if len(word) >= 2 and word not in stopwords:
                clean_word = ''.join(c for c in word if c.isalnum() or c in ['_', '-'])
                if clean_word and clean_word not in keywords:
                    keywords.append(clean_word)
                    if len(keywords) >= max_keywords:
                        break

        return keywords

    def _generate_chunk_contexts(
        self, chunks: List[str], full_text: str, filename: str
    ) -> List[str]:
        """각 청크에 대한 문맥 요약 생성 (Contextual Retrieval)

        Claude Haiku를 사용하여 전체 문서 내에서 각 청크의 위치와 맥락을
        50~100토큰 분량으로 요약합니다.

        Args:
            chunks: 청크 텍스트 리스트
            full_text: 문서 전체 텍스트
            filename: 파일명 (로깅용)

        Returns:
            각 청크에 대한 문맥 요약 리스트 (에러 시 빈 문자열)
        """
        try:
            from langchain_anthropic import ChatAnthropic

            llm = ChatAnthropic(
                model="claude-haiku-4-5",
                max_tokens=256,
                temperature=0.0,
            )
        except Exception as e:
            logger.warning(f"Claude Haiku 초기화 실패, 문맥 생성 건너뜀: {e}")
            return [""] * len(chunks)

        # 전체 문서가 너무 길면 앞부분만 사용 (토큰 제한)
        max_doc_chars = 15000
        doc_preview = full_text[:max_doc_chars]
        if len(full_text) > max_doc_chars:
            doc_preview += "\n...(이하 생략)"

        contexts = []
        for i, chunk in enumerate(chunks):
            try:
                prompt = (
                    f"<document>\n{doc_preview}\n</document>\n\n"
                    f"다음은 위 문서의 일부(청크)입니다:\n"
                    f"<chunk>\n{chunk}\n</chunk>\n\n"
                    f"이 청크가 문서 전체에서 어떤 맥락에 위치하는지 간결하게 설명하세요. "
                    f"50~100토큰 이내로, 한국어로 작성하세요. "
                    f"'이 청크는...' 같은 메타 표현 없이, "
                    f"문서의 주제와 이 부분의 핵심 내용만 서술하세요."
                )
                response = llm.invoke(prompt)
                context = response.content.strip() if response.content else ""
                contexts.append(context)
                # 생성된 문맥 전체를 로그로 출력 (최대 200자)
                context_preview = context[:200] + "..." if len(context) > 200 else context
                logger.info(f"[문맥 생성] {filename} 청크 {i+1}/{len(chunks)}:\n{context_preview}")
            except Exception as e:
                logger.warning(f"[문맥 생성] {filename} 청크 {i+1} 실패: {e}")
                contexts.append("")

        generated_count = sum(1 for c in contexts if c)
        logger.info(
            f"[문맥 생성] {filename}: {generated_count}/{len(chunks)}개 청크 문맥 생성 완료"
        )
        return contexts

    def _load_documents(self) -> List[Document]:
        """지식베이스에서 문서 로드 및 청킹"""
        documents = []

        if not self.knowledge_base_path.exists():
            logger.warning(f"지식베이스 경로가 존재하지 않습니다: {self.knowledge_base_path}")
            return documents

        # 문서 파싱 함수들
        def parse_text_file(file_path: Path) -> str:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"텍스트 파일 읽기 실패 ({file_path}): {e}")
                return ""

        def parse_pdf(file_path: Path) -> str:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                logger.warning("pypdf가 설치되지 않았습니다. PDF 파일은 건너뜁니다.")
                return ""
            except Exception as e:
                logger.error(f"PDF 파싱 실패 ({file_path}): {e}")
                return ""

        def parse_docx(file_path: Path) -> str:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                logger.warning("python-docx가 설치되지 않았습니다. DOCX 파일은 건너뜁니다.")
                return ""
            except Exception as e:
                logger.error(f"DOCX 파싱 실패 ({file_path}): {e}")
                return ""

        # 지원하는 파일 확장자 및 파서 매핑
        parsers = {
            '.txt': parse_text_file,
            '.md': parse_text_file,
            '.pdf': parse_pdf,
            '.docx': parse_docx,
        }

        # 모든 문서 파일 찾기
        for ext, parser_func in parsers.items():
            for file_path in self.knowledge_base_path.rglob(f"*{ext}"):
                try:
                    logger.info(f"문서 로드 중: {file_path.name}")

                    content = parser_func(file_path)
                    if not content.strip():
                        logger.warning(f"빈 문서: {file_path.name}")
                        continue

                    chunks = self.text_splitter.split_text(content)

                    # Contextual Retrieval: 각 청크에 문맥 요약 생성
                    contexts = self._generate_chunk_contexts(
                        chunks, content, file_path.name
                    )

                    for i, chunk in enumerate(chunks):
                        # 청크 위치 정보
                        if len(chunks) == 1:
                            position = "full"
                        elif i == 0:
                            position = "beginning"
                        elif i == len(chunks) - 1:
                            position = "end"
                        else:
                            position = "middle"

                        # 섹션 제목 추출
                        section_title = ""
                        chunk_lines = chunk.split('\n')
                        for line in chunk_lines[:3]:
                            line = line.strip()
                            if line.startswith('#'):
                                section_title = line.lstrip('#').strip()
                                break

                        # 키워드 추출
                        keywords = self._extract_keywords_from_text(chunk, max_keywords=5)

                        # 문맥이 있으면 page_content에 추가 (임베딩 품질 향상)
                        context = contexts[i] if i < len(contexts) else ""
                        if context:
                            enriched_content = f"{context}\n\n{chunk}"
                        else:
                            enriched_content = chunk

                        doc = Document(
                            page_content=enriched_content,
                            metadata={
                                'source': str(file_path),
                                'filename': file_path.name,
                                'extension': ext,
                                'chunk_index': i,
                                'total_chunks': len(chunks),
                                'position': position,
                                'chunk_length': len(chunk),
                                'section_title': section_title,
                                'keywords': ', '.join(keywords),
                                'original_content': chunk,
                                'context': context,  # 문맥 추가 (평가용)
                            }
                        )
                        documents.append(doc)

                    logger.info(f"✓ 문서 로드 완료: {file_path.name} ({len(chunks)}개 청크)")

                except Exception as e:
                    logger.error(f"문서 로드 실패 ({file_path}): {e}")
                    continue

        logger.info(f"총 {len(documents)}개 문서 청크 로드 완료")
        return documents

    def _check_embedding_dimension_match(self) -> bool:
        """기존 벡터 DB의 임베딩 차원이 현재 모델과 일치하는지 확인"""
        try:
            test_embedding = self.embeddings.embed_query("test")
            # numpy 배열이나 리스트를 안전하게 처리
            current_dim = int(len(test_embedding)) if hasattr(test_embedding, '__len__') else 0

            temp_store = Chroma(
                persist_directory=str(self.chroma_db_path),
                embedding_function=self.embeddings
            )
            collection = temp_store._collection
            count = collection.count()
            # numpy 배열을 정수로 변환
            count_int = int(count) if hasattr(count, '__iter__') else count
            if count_int == 0:
                return True  # 비어있으면 호환으로 간주

            # 기존 DB에서 하나 가져와서 차원 확인
            result = collection.peek(limit=1)
            embeddings_list = result.get('embeddings') if result else None
            has_embeddings = embeddings_list is not None and int(len(embeddings_list)) > 0
            if has_embeddings:
                first_embedding = embeddings_list[0]
                # numpy 배열이면 정수로 변환
                db_dim = int(len(first_embedding))
                current_dim_int = int(current_dim)

                if db_dim != current_dim_int:
                    logger.warning(
                        f"임베딩 차원 불일치! DB: {db_dim}, 현재 모델: {current_dim_int}. "
                        f"벡터 DB를 재구축합니다."
                    )
                    return False
            return True
        except Exception as e:
            logger.warning(f"임베딩 차원 확인 실패: {e}")
            return True  # 확인 실패 시 기존 DB 사용

    def _rebuild_vectorstore(self):
        """기존 벡터 DB를 백업 후 재구축"""
        import shutil
        from datetime import datetime

        if self.chroma_db_path.exists():
            # 백업 생성
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = Path(f"{self.chroma_db_path}.backup.{timestamp}")

            try:
                logger.info(f"벡터 DB 백업 생성 중: {backup_path}")
                shutil.copytree(self.chroma_db_path, backup_path)
                logger.info(f"✓ 백업 완료: {backup_path}")
            except Exception as e:
                logger.error(f"백업 생성 실패: {e}")
                raise RuntimeError(f"벡터 DB 백업 실패: {e}")

            # 기존 DB 삭제
            try:
                logger.info("기존 벡터 DB 삭제 중...")
                shutil.rmtree(str(self.chroma_db_path))
                logger.info("기존 벡터 DB 삭제 완료")
            except Exception as e:
                logger.error(f"벡터 DB 삭제 실패: {e}")
                logger.info("백업에서 복구 시도 중...")
                shutil.copytree(backup_path, self.chroma_db_path)
                raise RuntimeError(f"벡터 DB 삭제 실패, 복구됨: {e}")

            # 오래된 백업 정리
            self._cleanup_old_backups(keep_count=3)

    def _cleanup_old_backups(self, keep_count: int = 3):
        """오래된 백업 파일 정리"""
        import shutil

        parent_dir = self.chroma_db_path.parent
        backup_pattern = f"{self.chroma_db_path.name}.backup.*"

        backups = sorted(
            parent_dir.glob(backup_pattern),
            key=lambda p: p.stat().st_mtime,
            reverse=True
        )

        for old_backup in backups[keep_count:]:
            try:
                shutil.rmtree(old_backup)
                logger.info(f"오래된 백업 삭제: {old_backup}")
            except Exception as e:
                logger.warning(f"백업 삭제 실패: {old_backup} - {e}")

    def _build_vectorstore_if_needed(self) -> bool:
        """벡터 DB가 없으면 자동으로 구축, 차원 불일치 시 재구축"""
        if self._vectorstore is not None:
            return True

        # 기존 DB가 있으면 차원 일치 여부 확인
        if self.chroma_db_path.exists() and any(self.chroma_db_path.iterdir()):
            if self._check_embedding_dimension_match():
                return False  # 기존 DB 사용
            else:
                self._rebuild_vectorstore()  # 차원 불일치 → 재구축

        if not self.knowledge_base_path.exists():
            logger.warning(f"지식베이스 경로가 존재하지 않습니다: {self.knowledge_base_path}")
            return False

        # 문서 찾기
        has_documents = False
        for ext in ['.txt', '.md', '.pdf', '.docx']:
            if any(self.knowledge_base_path.rglob(f"*{ext}")):
                has_documents = True
                break

        if not has_documents:
            logger.warning("지식베이스에 문서가 없습니다. 벡터 DB를 구축할 수 없습니다.")
            return False

        # 벡터 DB 자동 구축
        logger.info("벡터 DB가 없습니다. 자동으로 구축을 시작합니다...")
        try:
            documents = self._load_documents()

            if not documents:
                logger.warning("로드할 문서가 없습니다.")
                return False

            # Chroma DB 생성 (cosine distance 사용)
            logger.info(f"벡터 DB 구축 중... ({len(documents)}개 문서 청크)")
            self._vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=str(self.chroma_db_path),
                collection_metadata={
                    "hnsw:space": "cosine",
                    "hnsw:construction_ef": 200,
                    "hnsw:search_ef": 100,
                    "hnsw:M": 32,
                }
            )

            logger.info(f"✓ 벡터 DB 구축 완료: {len(documents)}개 문서 (distance: cosine)")

            return True

        except Exception as e:
            logger.error(f"벡터 DB 구축 실패: {e}")
            return False

    @property
    def vectorstore(self) -> Optional[Chroma]:
        """벡터 스토어 지연 로딩 및 자동 구축"""
        if self._vectorstore is None and self.available:
            try:
                self._build_vectorstore_if_needed()

                if self.chroma_db_path.exists() and any(self.chroma_db_path.iterdir()):
                    if self._vectorstore is None:
                        self._vectorstore = Chroma(
                            persist_directory=str(self.chroma_db_path),
                            embedding_function=self.embeddings
                        )

                    # 진단: ChromaDB distance 함수 확인
                    try:
                        collection = self._vectorstore._collection
                        if collection and hasattr(collection, 'metadata') and collection.metadata:
                            metadata = collection.metadata
                            distance_function = metadata.get('hnsw:space', 'unknown')
                            logger.info(f"ChromaDB distance 함수: {distance_function}")
                        else:
                            logger.info("ChromaDB distance 함수: 확인 불가 (metadata 없음)")
                    except Exception as e:
                        logger.warning(f"Distance 함수 확인 실패: {e}")

                    logger.info("벡터 DB 로드 완료")
                else:
                    logger.warning("벡터 DB가 아직 구축되지 않았습니다.")
            except Exception as e:
                logger.error(f"벡터 스토어 로드 실패: {e}")

        return self._vectorstore

    def _normalize_query(self, query: str) -> str:
        """검색 쿼리 정규화"""
        import re
        normalized = re.sub(r'\s+', ' ', query)
        normalized = normalized.strip()
        return normalized

    def _tokenize(self, text: str) -> List[str]:
        """텍스트를 토큰으로 분할 (한국어/영어 지원)"""
        import re
        text = text.lower()
        tokens = re.findall(r'[가-힣a-z0-9]+', text)
        return tokens

    def _build_bm25_index(self) -> bool:
        """BM25 인덱스 빌드"""
        if not self.available:
            return False

        try:
            if self.vectorstore is not None:
                collection = self.vectorstore._collection
                results = collection.get(include=['documents', 'metadatas'])

                if not results or not results['documents']:
                    logger.warning("BM25 인덱스 구축 실패: 문서가 없습니다.")
                    return False

                self._bm25_documents = []
                for i, (doc_text, metadata) in enumerate(zip(results['documents'], results['metadatas'])):
                    doc = Document(page_content=doc_text, metadata=metadata or {})
                    self._bm25_documents.append(doc)

            else:
                self._bm25_documents = self._load_documents()

            if not self._bm25_documents:
                logger.warning("BM25 인덱스 구축 실패: 문서가 없습니다.")
                return False

            tokenized_corpus = [
                self._tokenize(doc.page_content)
                for doc in self._bm25_documents
            ]

            self._bm25_index = BM25Okapi(tokenized_corpus)
            logger.info(f"✓ BM25 인덱스 구축 완료: {len(self._bm25_documents)}개 문서")
            return True

        except Exception as e:
            logger.error(f"BM25 인덱스 구축 실패: {e}")
            return False

    @property
    def bm25_index(self) -> Optional['BM25Okapi']:
        """BM25 인덱스 지연 로딩"""
        if self._bm25_index is None and self.available:
            self._build_bm25_index()
        return self._bm25_index

    @property
    def graph_manager(self) -> Optional['Neo4jGraphManager']:
        """Neo4j 그래프 매니저 지연 로딩"""
        if self._graph_manager is None and self.available:
            try:
                from react_agent.knowledge_graph import get_graph_manager
                self._graph_manager = get_graph_manager()
                if self._graph_manager.available:
                    self._graph_manager.connect()
                    logger.info("Neo4j 그래프 매니저 연결 완료")
                else:
                    logger.warning("Neo4j 그래프 매니저를 사용할 수 없습니다.")
            except Exception as e:
                logger.warning(f"Neo4j 그래프 매니저 로딩 실패: {e}")
                self._graph_manager = None
        return self._graph_manager

    def search_documents(self, query: str, k: int = 4, similarity_threshold: float = 0.55, include_context: bool = False) -> List[Dict[str, Any]]:
        """
        문서 검색 (코사인 유사도 기반)

        Args:
            query: 검색 쿼리
            k: 반환할 문서 수
            similarity_threshold: 코사인 유사도 임계값
            include_context: 문맥을 포함할지 여부 (기본: False, 청크만 반환)

        Returns:
            관련 문서 리스트
        """
        if not self.available or self.vectorstore is None:
            return []

        self._check_kb_changed()

        cache_manager = get_cache_manager()
        cache_content = f"{query}|k={k}|threshold={similarity_threshold}"
        cached_result = cache_manager.get("rag", cache_content)
        if cached_result is not None:
            return cached_result

        try:
            normalized_query = self._normalize_query(query)
            if normalized_query != query:
                logger.info(f"[검색] 쿼리 정규화: '{query}' → '{normalized_query}'")
            else:
                logger.info(f"[검색] 쿼리: '{query}'")

            try:
                total_docs = self.vectorstore._collection.count()
                logger.info(f"지식베이스: 총 {total_docs}개 문서 청크")
            except Exception as e:
                logger.warning(f"지식베이스 문서 수 확인 실패: {e}")

            docs_with_scores = self.vectorstore.similarity_search_with_score(normalized_query, k=k * 3)
            logger.info(f"벡터 검색: {len(docs_with_scores)}개 결과")

            filtered_docs = []
            seen_keys = set()
            rejected_count = 0

            for idx, (doc, distance) in enumerate(docs_with_scores):
                if distance < 0:
                    similarity = 1.0
                elif distance > 2.0:
                    similarity = max(0.0, 1.0 - (distance / 2.0))
                else:
                    similarity = 1.0 - distance

                if similarity < similarity_threshold:
                    rejected_count += 1
                    continue

                source = doc.metadata.get('source', 'unknown')
                chunk_index = doc.metadata.get('chunk_index', 0)
                doc_key = (source, chunk_index)

                if doc_key in seen_keys:
                    continue

                seen_keys.add(doc_key)

                # 문맥 포함 여부에 따라 content 구성
                if include_context:
                    context = doc.metadata.get('context', '')
                    original = doc.metadata.get('original_content', doc.page_content)
                    content = f"{context}\n\n{original}" if context else original
                else:
                    content = doc.metadata.get('original_content', doc.page_content)

                filtered_docs.append({
                    'content': content,
                    'source': source,
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'chunk_index': chunk_index,
                    'similarity': float(similarity)
                })

                if len(filtered_docs) >= k:
                    break

            if not filtered_docs:
                logger.warning(
                    f"임계값 {similarity_threshold} 미만: "
                    f"{len(docs_with_scores)}개 결과 모두 제외됨"
                )
                cache_manager.set("rag", cache_content, [], ttl=3600)
                return []

            logger.info(f"필터링 완료: {len(filtered_docs)}개 선택, {rejected_count}개 제외")

            for idx, doc in enumerate(filtered_docs[:5]):
                logger.info(f"  #{idx+1}: {doc['filename']} (유사도: {doc['similarity']:.3f})")

            cache_manager.set("rag", cache_content, filtered_docs)
            return filtered_docs

        except Exception as e:
            logger.error(f"[검색] 검색 실패: {e}", exc_info=True)
            return []

    def search_documents_hybrid(
        self,
        query: str,
        k: int = 3,
        alpha: float = 0.5,
        similarity_threshold: float = 0.7
    ) -> List[Dict[str, Any]]:
        """
        하이브리드 검색 (BM25 + 벡터 검색, RRF 사용)

        Args:
            query: 검색 쿼리
            k: 반환할 문서 수
            alpha: 벡터 검색 가중치 (음수면 RRF 사용)
            similarity_threshold: 최소 유사도 임계값

        Returns:
            관련 문서 리스트
        """
        if not self.available:
            return []

        self._check_kb_changed()

        cache_manager = get_cache_manager()
        cache_content = f"hybrid|{query}|k={k}|alpha={alpha}|threshold={similarity_threshold}"
        cached_result = cache_manager.get("rag", cache_content)
        if cached_result is not None:
            return cached_result

        try:
            normalized_query = self._normalize_query(query)
            if normalized_query != query:
                logger.info(f"[하이브리드] 쿼리 정규화: '{query}' → '{normalized_query}'")
            else:
                logger.info(f"[하이브리드] 쿼리: '{query}'")

            try:
                total_docs = self.vectorstore._collection.count()
                logger.info(f"지식베이스: 총 {total_docs}개 문서 청크")
            except Exception as e:
                logger.warning(f"지식베이스 문서 수 확인 실패: {e}")

            # 1. 벡터 검색
            vector_results = {}
            if self.vectorstore is not None:
                vector_docs = self.vectorstore.similarity_search_with_score(normalized_query, k=k * 3)
                print(f"[VECTOR] {len(vector_docs)} results")

                # 진단: 실제 distance 값 확인
                if vector_docs:
                    print(f"[VECTOR] Top 3 (distance → similarity):")
                    for idx, (doc, distance) in enumerate(vector_docs[:3]):
                        filename = doc.metadata.get('filename', 'unknown')
                        similarity = 1.0 - distance
                        status = "HIGH" if similarity >= 0.7 else "MED" if similarity >= 0.5 else "LOW"
                        print(f"  - {filename[:60]}: dist={distance:.4f} → sim={similarity:.4f} [{status}]")

                for doc, distance in vector_docs:
                    doc_key = (doc.metadata.get('source', ''), doc.metadata.get('chunk_index', 0))
                    similarity = max(0.0, 1.0 - min(distance, 2.0))
                    vector_results[doc_key] = {
                        'doc': doc,
                        'score': similarity
                    }

            # 2. BM25 검색
            bm25_results = {}
            if self.bm25_index is not None and len(self._bm25_documents) > 0:
                tokenized_query = self._tokenize(normalized_query)
                bm25_scores = self.bm25_index.get_scores(tokenized_query)

                max_score = max(bm25_scores) if len(bm25_scores) > 0 and max(bm25_scores) > 0 else 1.0
                normalized_scores = [score / max_score for score in bm25_scores]

                top_indices = np.argsort(normalized_scores)[::-1][:k * 3]
                print(f"[BM25] {len(top_indices)} results")

                for idx in top_indices:
                    doc = self._bm25_documents[idx]
                    doc_key = (doc.metadata.get('source', ''), doc.metadata.get('chunk_index', 0))
                    bm25_results[doc_key] = {
                        'doc': doc,
                        'score': normalized_scores[idx]
                    }

            # 3. 점수 결합 (RRF 또는 가중 평균)
            combined_results = {}
            all_doc_keys = set(vector_results.keys()) | set(bm25_results.keys())

            use_rrf = alpha < 0  # 음수 alpha는 RRF 활성화 신호

            if use_rrf:
                # RRF 방식: 순위로 정렬, 점수는 vector+bm25 평균 사용
                k_rrf = 60

                # Vector 결과 순위 생성
                vector_sorted = sorted(
                    vector_results.items(),
                    key=lambda x: x[1]['score'],
                    reverse=True
                )
                vector_ranks = {doc_key: rank for rank, (doc_key, _) in enumerate(vector_sorted)}

                # BM25 결과 순위 생성
                bm25_sorted = sorted(
                    bm25_results.items(),
                    key=lambda x: x[1]['score'],
                    reverse=True
                )
                bm25_ranks = {doc_key: rank for rank, (doc_key, _) in enumerate(bm25_sorted)}

                # RRF로 순위 계산 (Numpy vectorization으로 최적화)
                all_doc_keys_list = list(all_doc_keys)

                # Vectorized rank extraction
                vector_ranks_array = np.array([
                    vector_ranks.get(doc_key, len(vector_results))
                    for doc_key in all_doc_keys_list
                ])
                bm25_ranks_array = np.array([
                    bm25_ranks.get(doc_key, len(bm25_results))
                    for doc_key in all_doc_keys_list
                ])

                # Vectorized RRF calculation: 1/(k + rank + 1)
                rrf_rank_scores = (1.0 / (k_rrf + vector_ranks_array + 1) +
                                   1.0 / (k_rrf + bm25_ranks_array + 1))

                # Vectorized score extraction
                vector_scores = np.array([
                    vector_results.get(doc_key, {}).get('score', 0.0)
                    for doc_key in all_doc_keys_list
                ])
                bm25_scores = np.array([
                    bm25_results.get(doc_key, {}).get('score', 0.0)
                    for doc_key in all_doc_keys_list
                ])

                # 최종 점수: vector와 bm25의 평균 (절대적 품질 유지)
                hybrid_scores = (vector_scores + bm25_scores) / 2.0

                # 결과 저장
                for idx, doc_key in enumerate(all_doc_keys_list):
                    doc = vector_results.get(doc_key, bm25_results.get(doc_key, {})).get('doc')

                    if doc is not None:
                        combined_results[doc_key] = {
                            'doc': doc,
                            'rrf_rank_score': float(rrf_rank_scores[idx]),  # 정렬용
                            'hybrid_score': float(hybrid_scores[idx]),  # 임계값 필터링용
                            'vector_score': float(vector_scores[idx]),
                            'bm25_score': float(bm25_scores[idx])
                        }
            else:
                # 가중 평균 방식
                for doc_key in all_doc_keys:
                    vector_score = vector_results.get(doc_key, {}).get('score', 0.0)
                    bm25_score = bm25_results.get(doc_key, {}).get('score', 0.0)

                    hybrid_score = alpha * vector_score + (1.0 - alpha) * bm25_score

                    doc = vector_results.get(doc_key, bm25_results.get(doc_key, {})).get('doc')

                    if doc is not None:
                        combined_results[doc_key] = {
                            'doc': doc,
                            'hybrid_score': hybrid_score,
                            'vector_score': vector_score,
                            'bm25_score': bm25_score
                        }

            # 4. 하이브리드 점수로 정렬
            if use_rrf:
                sorted_results = sorted(
                    combined_results.items(),
                    key=lambda x: x[1].get('rrf_rank_score', x[1]['hybrid_score']),
                    reverse=True
                )
            else:
                sorted_results = sorted(
                    combined_results.items(),
                    key=lambda x: x[1]['hybrid_score'],
                    reverse=True
                )

            # 5. 필터링 및 결과 생성
            filtered_docs = []
            rejected_count = 0
            for idx, (doc_key, result) in enumerate(sorted_results):
                hybrid_score = result['hybrid_score']

                if hybrid_score < similarity_threshold:
                    rejected_count += 1
                    continue

                doc = result['doc']
                filtered_docs.append({
                    'content': doc.metadata.get('original_content', doc.page_content),
                    'source': doc.metadata.get('source', 'unknown'),
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'chunk_index': doc.metadata.get('chunk_index', 0),
                    'similarity': float(hybrid_score),
                    'vector_score': float(result['vector_score']),
                    'bm25_score': float(result['bm25_score'])
                })

                if len(filtered_docs) >= k:
                    break

            if not filtered_docs:
                logger.warning(
                    f"임계값 {similarity_threshold} 미만: "
                    f"{len(sorted_results)}개 결과 모두 제외됨"
                )
                cache_manager.set("rag", cache_content, [], ttl=3600)
                return []

            logger.info(f"필터링 완료: {len(filtered_docs)}개 선택, {rejected_count}개 제외")

            logger.info(f"최종 결과 (하이브리드, alpha={alpha}): {len(filtered_docs)}개")
            for idx, doc in enumerate(filtered_docs[:5]):
                logger.info(
                    f"  #{idx+1}: {doc['filename']} "
                    f"(hybrid: {doc['similarity']:.3f} = vector: {doc['vector_score']:.3f} + bm25: {doc['bm25_score']:.3f})"
                )

            cache_manager.set("rag", cache_content, filtered_docs)
            return filtered_docs

        except Exception as e:
            logger.error(f"[하이브리드] 검색 실패: {e}", exc_info=True)
            return []

    def search_documents_graph(
        self,
        query: str,
        k: int = 3,
        similarity_threshold: float = 0.7,
        graph_boost: float = 0.3
    ) -> List[Dict[str, Any]]:
        """
        그래프 기반 하이브리드 검색 (벡터 + Neo4j 주제 확장)

        알고리즘:
        1. 벡터 검색으로 top k*3 후보 조회
        2. Neo4j에서 주제/인접 청크 확장
        3. 최종 점수 = vector_score + graph_boost (cap 1.0)
        4. threshold 필터링 → top-k 반환

        Args:
            query: 검색 쿼리
            k: 반환할 문서 수
            similarity_threshold: 최소 유사도 임계값
            graph_boost: 그래프 연결 청크의 추가 점수 (부모 점수 × boost)

        Returns:
            관련 문서 리스트
        """
        if not self.available or self.vectorstore is None:
            return []

        self._check_kb_changed()

        # 캐시 확인
        cache_manager = get_cache_manager()
        cache_content = f"graph|{query}|k={k}|threshold={similarity_threshold}|boost={graph_boost}"
        cached_result = cache_manager.get("rag", cache_content)
        if cached_result is not None:
            return cached_result

        try:
            t_total_start = time.perf_counter()
            normalized_query = self._normalize_query(query)
            logger.info(f"[그래프 검색] 쿼리: '{normalized_query}'")

            # 1. 벡터 검색: top k*3 후보
            t_vector_start = time.perf_counter()
            logger.debug(f"[임베딩 입력] \"{normalized_query}\"")
            vector_docs = self.vectorstore.similarity_search_with_score(
                normalized_query, k=k * 3
            )
            vector_elapsed = time.perf_counter() - t_vector_start

            if not vector_docs:
                logger.warning("[그래프 검색] 벡터 검색 결과 없음")
                cache_manager.set("rag", cache_content, [], ttl=3600)
                return []

            logger.info(f"[그래프 검색] 벡터 후보: {len(vector_docs)}개")
            perf_logger.info(f"⏱️ [그래프 검색 > 벡터] {vector_elapsed:.2f}초 (쿼리: \"{normalized_query[:50]}...\", 후보: {len(vector_docs)})")

            # 벡터 결과를 dict로 저장 (점수 포함)
            vector_results: Dict[Tuple[str, int], Dict[str, Any]] = {}
            for doc, distance in vector_docs:
                filename = doc.metadata.get('filename', '')
                source = doc.metadata.get('source', '')
                chunk_index = doc.metadata.get('chunk_index', 0)
                doc_key = (source, chunk_index)

                # cosine distance → similarity
                similarity = max(0.0, 1.0 - min(distance, 2.0))

                vector_results[doc_key] = {
                    'doc': doc,
                    'filename': filename,
                    'source': source,
                    'chunk_index': chunk_index,
                    'vector_score': similarity,
                    'graph_boost': 0.0,
                    'is_graph_expanded': False
                }

            # 2. 그래프 확장
            t_graph_start = time.perf_counter()
            graph_manager = self.graph_manager
            expanded_chunks: Dict[Tuple[str, int], Dict[str, Any]] = {}

            if graph_manager and graph_manager.available:
                # 벡터 결과의 청크 ID 추출 (filename으로 조회)
                chunk_ids = [
                    (v['filename'], v['chunk_index'])
                    for v in vector_results.values()
                ]

                # Neo4j에서 관련 청크 조회
                related_map = graph_manager.get_related_chunks(chunk_ids)

                for parent_key, related in related_map.items():
                    # parent_key는 (filename, chunk_index)
                    # vector_results의 key는 (source, chunk_index)
                    # 매핑 필요: filename으로 source 찾기
                    parent_result = None
                    for vk, vv in vector_results.items():
                        if vv['filename'] == parent_key[0] and vv['chunk_index'] == parent_key[1]:
                            parent_result = vv
                            break

                    if parent_result is None:
                        continue

                    parent_score = parent_result['vector_score']

                    # 개념 연결 청크 (주요 확장)
                    for filename, chunk_index, shared_concepts in related.get('concept_related', []):
                        expanded_key = (filename, chunk_index)

                        if expanded_key in vector_results:
                            # 이미 벡터 결과에 있으면 boost만 추가
                            # 동일 source로 매핑
                            for vk, vv in vector_results.items():
                                if vv['filename'] == filename and vv['chunk_index'] == chunk_index:
                                    current_boost = vv['graph_boost']
                                    new_boost = parent_score * graph_boost
                                    vv['graph_boost'] = max(current_boost, new_boost)
                                    vv['is_graph_expanded'] = True
                                    break
                        else:
                            # 새로운 청크 추가
                            boost_score = parent_score * graph_boost

                            if expanded_key not in expanded_chunks:
                                expanded_chunks[expanded_key] = {
                                    'filename': filename,
                                    'chunk_index': chunk_index,
                                    'vector_score': 0.0,  # 벡터 결과에 없음
                                    'graph_boost': boost_score,
                                    'is_graph_expanded': True,
                                    'expansion_concepts': shared_concepts,
                                    'parent_chunk': parent_key
                                }
                            else:
                                # 이미 다른 부모에서 확장됨 → boost 최대값 유지
                                expanded_chunks[expanded_key]['graph_boost'] = max(
                                    expanded_chunks[expanded_key]['graph_boost'],
                                    boost_score
                                )

                    # 인접 청크 (보조 확장, 낮은 boost)
                    for filename, chunk_index in related.get('neighbors', []):
                        neighbor_key = (filename, chunk_index)
                        neighbor_boost = parent_score * graph_boost * 0.5  # 인접은 절반

                        if neighbor_key in vector_results:
                            for vk, vv in vector_results.items():
                                if vv['filename'] == filename and vv['chunk_index'] == chunk_index:
                                    # 주제 확장보다 인접 확장이 우선되지 않도록 max 사용
                                    if not vv['is_graph_expanded']:
                                        vv['graph_boost'] = max(vv['graph_boost'], neighbor_boost)
                                        vv['is_graph_expanded'] = True
                                    break
                        elif neighbor_key not in expanded_chunks:
                            expanded_chunks[neighbor_key] = {
                                'filename': filename,
                                'chunk_index': chunk_index,
                                'vector_score': 0.0,
                                'graph_boost': neighbor_boost,
                                'is_graph_expanded': True,
                                'expansion_topic': 'neighbor',
                                'parent_chunk': parent_key
                            }

                graph_elapsed = time.perf_counter() - t_graph_start
                logger.info(f"[그래프 검색] 그래프 확장: {len(expanded_chunks)}개 추가 청크")
                perf_logger.info(f"⏱️ [그래프 검색 > 그래프 확장] {graph_elapsed:.2f}초 (확장: {len(expanded_chunks)})")

            # 3. 확장된 청크의 문서 내용 로드 (벡터 DB에서 조회)
            if expanded_chunks and self.vectorstore is not None:
                collection = self.vectorstore._collection
                for expanded_key, expanded_info in expanded_chunks.items():
                    filename = expanded_info['filename']
                    chunk_index = expanded_info['chunk_index']

                    # ChromaDB에서 해당 청크 조회
                    try:
                        results = collection.get(
                            where={
                                "$and": [
                                    {"filename": {"$eq": filename}},
                                    {"chunk_index": {"$eq": chunk_index}}
                                ]
                            },
                            include=['documents', 'metadatas']
                        )

                        if results and results['documents'] and len(results['documents']) > 0:
                            doc_text = results['documents'][0]
                            metadata = results['metadatas'][0] if results['metadatas'] else {}
                            expanded_info['doc'] = Document(
                                page_content=doc_text,
                                metadata=metadata
                            )
                            expanded_info['source'] = metadata.get('source', '')
                    except Exception as e:
                        logger.warning(f"확장 청크 조회 실패 ({filename}:{chunk_index}): {e}")

            # 4. 모든 결과 통합 및 점수 계산
            all_results: Dict[Tuple[str, int], Dict[str, Any]] = {}

            # 벡터 결과 추가
            for key, result in vector_results.items():
                final_score = result['vector_score'] + result['graph_boost']
                final_score = min(final_score, 1.0)  # cap at 1.0
                all_results[key] = {
                    **result,
                    'final_score': final_score
                }

            # 확장된 청크 추가 (문서가 로드된 것만)
            for key, result in expanded_chunks.items():
                if 'doc' in result:
                    final_score = result['vector_score'] + result['graph_boost']
                    final_score = min(final_score, 1.0)
                    # source로 key 변환
                    source_key = (result.get('source', key[0]), key[1])
                    if source_key not in all_results:
                        all_results[source_key] = {
                            **result,
                            'final_score': final_score
                        }

            # 5. 점수 기준 정렬
            sorted_results = sorted(
                all_results.items(),
                key=lambda x: x[1]['final_score'],
                reverse=True
            )

            # 6. 필터링 및 결과 생성
            filtered_docs = []
            rejected_count = 0

            for doc_key, result in sorted_results:
                final_score = result['final_score']

                if final_score < similarity_threshold:
                    rejected_count += 1
                    continue

                doc = result.get('doc')
                if doc is None:
                    continue

                filtered_docs.append({
                    'content': doc.metadata.get('original_content', doc.page_content),
                    'source': doc.metadata.get('source', result.get('source', 'unknown')),
                    'filename': doc.metadata.get('filename', result.get('filename', 'unknown')),
                    'chunk_index': doc.metadata.get('chunk_index', result.get('chunk_index', 0)),
                    'similarity': float(final_score),
                    'vector_score': float(result['vector_score']),
                    'graph_boost': float(result['graph_boost']),
                    'is_graph_expanded': result.get('is_graph_expanded', False)
                })

                if len(filtered_docs) >= k:
                    break

            if not filtered_docs:
                logger.warning(
                    f"[그래프 검색] 임계값 {similarity_threshold} 미만: "
                    f"{len(sorted_results)}개 결과 모두 제외됨"
                )
                cache_manager.set("rag", cache_content, [], ttl=3600)
                return []

            logger.info(f"[그래프 검색] 결과: {len(filtered_docs)}개 선택, {rejected_count}개 제외")

            for idx, doc in enumerate(filtered_docs[:5]):
                graph_tag = " [GRAPH]" if doc['is_graph_expanded'] else ""
                logger.info(
                    f"  #{idx+1}: {doc['filename']} "
                    f"(final: {doc['similarity']:.3f} = vec: {doc['vector_score']:.3f} + boost: {doc['graph_boost']:.3f}){graph_tag}"
                )

            total_elapsed = time.perf_counter() - t_total_start
            perf_logger.info(f"⏱️ [그래프 검색] 총 {total_elapsed:.2f}초 (결과: {len(filtered_docs)}건)")

            cache_manager.set("rag", cache_content, filtered_docs)
            return filtered_docs

        except Exception as e:
            logger.error(f"[그래프 검색] 검색 실패: {e}", exc_info=True)
            return []

    async def build_knowledge_graph(self) -> bool:
        """지식 그래프 구축 (인덱싱 시 호출)

        ChromaDB의 문서들을 Neo4j 그래프로 변환합니다.
        각 청크에서 LLM으로 개념을 추출하고 도메인 기반 그래프를 구축합니다.
        - Domain 노드: 금융, 공공, 의료, 법률, 상업
        - Chunk 노드: 개념 정보 포함
        - RELATED_TO 관계: 공유 개념 기반 청크 연결

        Returns:
            성공 여부
        """
        graph_manager = self.graph_manager
        if graph_manager is None or not graph_manager.available:
            logger.warning("Neo4j 그래프 매니저를 사용할 수 없습니다.")
            return False

        if self.vectorstore is None:
            logger.warning("벡터 스토어가 없습니다. 먼저 벡터 DB를 구축하세요.")
            return False

        try:
            # ChromaDB에서 모든 문서 조회
            collection = self.vectorstore._collection
            results = collection.get(include=['documents', 'metadatas'])

            if not results or not results['documents']:
                logger.warning("ChromaDB에 문서가 없습니다.")
                return False

            # Document 형식으로 변환
            documents = []
            for doc_text, metadata in zip(results['documents'], results['metadatas'] or [{}] * len(results['documents'])):
                documents.append({
                    'page_content': doc_text,
                    'metadata': metadata or {}
                })

            logger.info(f"지식 그래프 구축 시작: {len(documents)}개 문서")
            print(f"[DEBUG] 지식 그래프 구축 시작: {len(documents)}개 문서")

            # 디버깅: 첫 문서 확인
            if documents:
                logger.info(f"첫 문서 샘플: metadata={documents[0].get('metadata', {})}")
                print(f"[DEBUG] 첫 문서 샘플: metadata={documents[0].get('metadata', {})}")
            else:
                logger.warning("문서가 비어있습니다!")
                print(f"[DEBUG] 문서가 비어있습니다!")

            # 기존 그래프 초기화 (선택적)
            # graph_manager.clear_graph()

            # 그래프 구축
            success = await graph_manager.build_graph(documents)

            if success:
                stats = graph_manager.get_graph_stats()
                logger.info(f"✓ 지식 그래프 구축 완료: {stats}")

            return success

        except Exception as e:
            logger.error(f"지식 그래프 구축 실패: {e}", exc_info=True)
            return False


# 전역 RAG 도구 인스턴스
_rag_tool: Optional[RAGTool] = None


def get_rag_tool() -> RAGTool:
    """RAG 도구 싱글톤 인스턴스 반환"""
    global _rag_tool
    if _rag_tool is None:
        _rag_tool = RAGTool()
    return _rag_tool
